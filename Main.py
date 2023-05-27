from docx import Document
from docx.shared import Inches
import subprocess
import os, sys
if getattr(sys, 'frozen', False):
    program_directory = os.path.dirname(os.path.abspath(sys.executable))
else:
    program_directory = os.path.dirname(os.path.abspath(__file__))


def asExec(ascript):
    osa = subprocess.Popen(['osascript', '-'],
                           stdin=subprocess.PIPE,
                           stdout=subprocess.PIPE)
    return osa.communicate(ascript)[0]
def asConv(astr):
    astr = astr.replace('"', '" & quote & "')
    return '"{}"'.format(astr)
def aScript(aspath):
    ascript = '''
    set posixPath to POSIX path of {0}
    '''.format(asConv(aspath))
    return ascript.encode()

str1= program_directory + "/AltText1.txt"
str2= program_directory + "/AltTextpath.txt"


with open(str2) as f1:
    #contents2 = f1.readlines()
    writefile = [x.strip() for x in f1.readlines()]

#txt44=asExec(aScript(writefile[0])).decode().strip()

with open(str1) as f:
    lines = [x.strip() for x in f.readlines()]
document = Docuament()

for line in lines:
    tmp = line.strip().lower()
    split_line = tmp.split(';')
    #print(split_line[2])
    txt1="Image Name: "+split_line[0]
    txt2=asExec(aScript(split_line[1])).decode().strip()
    
    #resl=(asExec(aScript((txt2))))
    #print(resl.decode())
    txt3="Alt Text: "+split_line[2]
    

    p = document.add_paragraph()
    r = p.add_run()
    r.add_text("----------------------------------------------------\n")
    r.add_text(txt1)
    r.add_text("\n")
    r.add_picture(txt2)
    r.add_text("\n")
    #r.add_picture(str1)
    if split_line[2]!="":
        r.add_text(txt3)
    #print(split_line[1])
    #print(asExec(aScript(split_line[1])).decode().strip())

    #if split_line[0]=='bike' and int(split_line[1])>0:
      #  bike_available = True


#aliasPath = "Main HD:Users:sasha:Documents:SomeText.txt"
#print(asExec(aScript(aliasPath)))

#print(asExec(aScript(split_line[1])))
#print(writefile[0])
document.save(writefile[0])
